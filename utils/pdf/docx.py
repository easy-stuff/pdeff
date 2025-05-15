import io
import logging
import os
import tempfile
import zipfile

from docx import Document
from docx.shared import Inches, Pt
from docx2pdf import convert
from pdf2docx import Converter
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
import win32com.client
import pythoncom 

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)


def pixels_to_docx_units(pixels, dpi=300):
    inches = pixels / dpi
    return Inches(inches)

def pt_to_emu(pt):
        # 1 pt = 1/72 inch, 1 inch = 914400 EMUs => 1 pt = 914400/72 = ~12700
        return int(pt * 12700)

def to_docx_no_ocr(files: list) -> io.BytesIO:
    logger.debug("Starting PDF to DOCX conversion using pdf2docx with page size adjustments")

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        with tempfile.TemporaryDirectory() as tmpdir:
            for index, file in enumerate(files):
                try:
                    filename = getattr(file, 'filename', f'file_{index}.pdf')
                    base_name = os.path.splitext(filename)[0]
                    logger.debug(f"Processing file: {filename}")

                    file_path = os.path.join(tmpdir, filename)
                    file.seek(0)
                    with open(file_path, 'wb') as f:
                        f.write(file.read())
                    logger.debug(f"Saved PDF to: {file_path}")

                    docx_filename = f"{base_name}.docx"
                    raw_docx_path = os.path.join(tmpdir, f"raw_{docx_filename}")
                    final_docx_path = os.path.join(tmpdir, docx_filename)

                    cv = Converter(file_path)
                    cv.convert(raw_docx_path, start=0, end=None)
                    cv.close()
                    logger.debug(f"Converted to DOCX: {raw_docx_path}")

                    pdf_doc = fitz.open(file_path)
                    doc = Document(raw_docx_path)

                    for i, page in enumerate(pdf_doc):
                        width_pt, height_pt = page.rect.width, page.rect.height
                        if i == 0:
                            section = doc.sections[0]
                        else:
                            section = doc.add_section()
                        section.page_width = pt_to_emu(width_pt)
                        section.page_height = pt_to_emu(height_pt)

                        logger.debug(f"Page {i + 1}: set size {width_pt:.2f}pt x {height_pt:.2f}pt")

                    pdf_doc.close()
                    doc.save(final_docx_path)

                    # Step 3: Add to ZIP
                    with open(final_docx_path, 'rb') as docx_file:
                        zip_file.writestr(docx_filename, docx_file.read())
                        logger.debug(f"Added {docx_filename} to ZIP")

                except Exception as e:
                    logger.exception(f"Error processing {filename}: {e}")

    zip_buffer.seek(0)
    logger.debug("Returning zipped DOCX files as BytesIO")
    return zip_buffer

def to_docx_ocr(files: list) -> io.BytesIO:
    logger.debug("Starting OCR-based PDF to DOCX conversion with page size adjustment")

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        with tempfile.TemporaryDirectory() as tmpdir:
            for index, file in enumerate(files):
                try:
                    filename = getattr(file, 'filename', f'file_{index}.pdf')
                    base_name = os.path.splitext(filename)[0]
                    logger.debug(f"OCR processing: {filename}")

                    # Save PDF to temp location
                    pdf_path = os.path.join(tmpdir, filename)
                    file.seek(0)
                    with open(pdf_path, 'wb') as f:
                        f.write(file.read())

                    # Convert PDF to images
                    images = convert_from_path(pdf_path, dpi=300)
                    logger.debug(f"Converted {len(images)} pages to images")

                    doc = Document()

                    for i, image in enumerate(images):
                        # Set section size based on image size
                        width, height = image.size  # in pixels
                        if i == 0:
                            section = doc.sections[0]
                        else:
                            section = doc.add_section()
                        section.page_width = pixels_to_docx_units(width)
                        section.page_height = pixels_to_docx_units(height)

                        # OCR text
                        text = pytesseract.image_to_string(image)
                        doc.add_paragraph(text)

                        logger.debug(f"OCR done for page {i + 1}, size {width}x{height}px")

                    # Save to docx
                    docx_filename = f"{base_name}_ocr.docx"
                    docx_path = os.path.join(tmpdir, docx_filename)
                    doc.save(docx_path)

                    # Add to zip
                    with open(docx_path, 'rb') as f:
                        zip_file.writestr(docx_filename, f.read())
                        logger.debug(f"Added {docx_filename} to zip")

                except Exception as e:
                    logger.exception(f"Error processing {filename}: {e}")

    zip_buffer.seek(0)
    logger.debug("Returning final ZIP file")
    return zip_buffer

def to_pdf(files: list) -> io.BytesIO:
    logger.debug("Starting DOC/DOCX to PDF conversion")

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        with tempfile.TemporaryDirectory() as tmpdir:
            for index, file in enumerate(files):
                try:
                    filename = getattr(file, 'filename', f'file_{index}.docx')
                    base_name = os.path.splitext(filename)[0]
                    logger.debug(f"Processing file: {filename}")

                    input_path = os.path.join(tmpdir, filename)
                    file.seek(0)
                    with open(input_path, 'wb') as f:
                        f.write(file.read())
                    logger.debug(f"Saved input to: {input_path}")

                    output_path = os.path.join(tmpdir, base_name + ".pdf")

                    pythoncom.CoInitialize()
                    try:
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        wdFormatPDF = 17

                        doc = word.Documents.Open(str(input_path))
                        doc.SaveAs(str(output_path), FileFormat=wdFormatPDF)
                        doc.Close(False)
                        word.Quit()
                        logger.debug(f"Converted to PDF: {output_path}")
                    finally:
                        pythoncom.CoUninitialize()

                    if not os.path.exists(output_path):
                        raise FileNotFoundError(f"PDF not created for {filename}")

                    with open(output_path, 'rb') as pdf_file:
                        zip_file.writestr(base_name + ".pdf", pdf_file.read())
                        logger.debug(f"Added {base_name}.pdf to ZIP")

                except Exception as e:
                    logger.exception(f"Error handling {filename}: {e}")

    zip_buffer.seek(0)
    logger.debug("Returning zipped PDF files as BytesIO")
    return zip_buffer
